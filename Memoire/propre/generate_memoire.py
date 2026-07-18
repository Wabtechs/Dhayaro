from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = 'Times New Roman'
    heading_font.color.rgb = RGBColor(0, 0, 0)
    heading_font.bold = True
    if i == 1:
        heading_font.size = Pt(16)
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif i == 2:
        heading_font.size = Pt(14)
    elif i == 3:
        heading_font.size = Pt(13)
    else:
        heading_font.size = Pt(12)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("UNIVERSITE DE BUKAVU")
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

fac = doc.add_paragraph()
fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fac.add_run("FACULTÉ DES SCIENCES APPLIQUÉES")
run.font.size = Pt(13)
run.font.bold = True
run.font.name = 'Times New Roman'

dept = doc.add_paragraph()
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dept.add_run("Département d'Informatique")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

main_title = doc.add_paragraph()
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = main_title.add_run("MODÉLISATION ET IMPLÉMENTATION D'UN ÉCOSYSTÈME NUMÉRIQUE RÉSILIENT POUR LA GESTION HOSPITALIÈRE ET L'ARCHIVAGE PÉRENNE DES PROTOCOLES DE SOINS EN ZONES DÉCENTRALISÉES")
run.font.size = Pt(15)
run.font.bold = True
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("CAS DE LA ZONE DE SANTÉ D'ALIMBONGO/NORD-KIVU")
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Mémoire présenté en vue de l'obtention du titre de Licencié en Informatique")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run("Année Académique 2025-2026")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

doc.add_heading("DÉDICACE", level=1)
doc.add_paragraph("Je dédie ce travail à mes chers parents, à mes frères et sœurs, et à tous ceux qui ont cru en moi et m'ont soutenu tout au long de ce parcours académique.")

doc.add_heading("REMERCIEMENTS", level=1)
doc.add_paragraph("Nous tenons à exprimer notre profonde gratitude à tous ceux qui, de près ou de loin, ont contribué à la réalisation de ce travail. Nous remercions en particulier notre directeur de mémoire pour sa guidance, les membres du corps enseignant de la Faculté des Sciences Appliquées de l'Université de Bukavu, ainsi que le personnel de la Zone de Santé d'Alimbongo pour leur collaboration et leur accueil chaleureux.")

doc.add_heading("LISTE DES ABRÉVIATIONS", level=1)

abbreviations = [
    ("BCZS", "Bureau Central de la Zone de Santé"),
    ("DMP", "Dossier Médical Partagé"),
    ("DPI", "Dossier Patient Informatisé"),
    ("DPS", "Divisions Provinciales de la Santé"),
    ("DHIS2", "District Health Information Software 2"),
    ("FHIR", "Fast Healthcare Interoperability Resources"),
    ("HGR", "Hôpital Général de Référence"),
    ("HL7", "Health Level Seven"),
    ("OMS", "Organisation Mondiale de la Santé"),
    ("OMOP CDM", "Observational Medical Outcomes Partnership Common Data Model"),
    ("PDF/A", "Portable Document Format for Archiving"),
    ("SNIS", "Système National d'Information Sanitaire"),
    ("SIH", "Système d'Information Hospitalier"),
    ("ZS", "Zone de Santé"),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Abréviation'
hdr_cells[1].text = 'Signification'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)

for abbr, full in abbreviations:
    row_cells = table.add_row().cells
    row_cells[0].text = abbr
    row_cells[1].text = full
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)

doc.add_page_break()

doc.add_heading("INTRODUCTION GÉNÉRALE", level=1)

doc.add_heading("1. État de la question (Contexte global)", level=2)
doc.add_paragraph("À l'ère de la transformation numérique, la gestion des systèmes de santé mondiale repose de plus en plus sur l'informatique pour améliorer la qualité des soins et la sécurité des données. Dans les pays développés, les Dossiers Médicaux Partagés (DMP) et les protocoles de soins numérisés sont devenus la norme, permettant une prise en charge rapide et un archivage sécurisé.")

doc.add_heading("2. Problématique (Le contraste local)", level=2)
doc.add_paragraph("Cependant, dans les pays en développement, et particulièrement en République Démocratique du Congo, l'accès à ces technologies se heurte à des réalités géographiques et infrastructurelles complexes. La Zone de Santé d'Alimbongo, située dans une région décentralisée et souvent enclavée, illustre parfaitement ce défi. Le système de santé y reste dominé par une gestion manuelle sur papier, entraînant une fragilité des archives (pertes, usures), une lenteur dans la transmission des rapports épidémiologiques et une difficulté pour le personnel soignant d'accéder en temps réel aux protocoles de soins standardisés. À cela s'ajoutent l'instabilité du réseau électrique et la rareté de la connexion internet, qui rendent les solutions numériques classiques inadaptées.")

doc.add_paragraph("Dès lors, une question fondamentale se pose : Comment concevoir et implémenter un écosystème numérique qui soit à la fois performant pour la gestion hospitalière et assez résilient pour survivre aux contraintes techniques d'une zone décentralisée comme Alimbongo ?")

doc.add_heading("3. Hypothèse", level=2)
doc.add_paragraph("Pour répondre à cette problématique, nous soutenons qu'un écosystème numérique hybride (Offline-first), fonctionnant sur des infrastructures à basse consommation énergétique et intégrant des protocoles de soins interactifs, permettrait de garantir la continuité du service et la pérennité des archives médicales, indépendamment des aléas de connexion et d'énergie.")

doc.add_heading("4. Objectifs", level=2)
doc.add_paragraph("L'objectif général de ce travail est de modéliser et de mettre en œuvre une solution informatique résiliente capable d'automatiser les processus hospitaliers et d'assurer un archivage durable des données cliniques à Alimbongo. Plus spécifiquement, il s'agira de :")

objectives = [
    "Concevoir une architecture de base de données capable de supporter des synchronisations intermittentes.",
    "Digitaliser les protocoles de soins pour l'aide à la décision clinique.",
    "Assurer la sécurité et la redondance des données en milieu rural.",
    "Optimiser le système pour qu'il puisse fonctionner sur des équipements à très faible consommation d'énergie, adaptés aux kits solaires disponibles à Alimbongo."
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading("5. Choix et intérêt du sujet", level=2)
doc.add_paragraph("Intérêt scientifique : Contribuer à la recherche sur les systèmes d'information en zones à ressources limitées (Edge Computing).")
doc.add_paragraph("Intérêt social : Améliorer la prise en charge des patients dans la zone de santé d'Alimbongo par une meilleure gestion des informations médicales.")

doc.add_heading("6. Délimitation du sujet", level=2)
doc.add_paragraph("Ce travail se focalise sur la Zone de Santé d'Alimbongo, dans la province du Nord-Kivu, et couvre la période allant de 2024 à 2026.")

doc.add_heading("7. Subdivision du travail", level=2)
doc.add_paragraph("Hormis l'introduction et la conclusion, ce travail s'articule autour de trois chapitres :")
chapters = [
    "Chapitre 1 : Considérations théoriques et définition des concepts.",
    "Chapitre 2 : Analyse de l'existant et étude des besoins (Cas d'Alimbongo).",
    "Chapitre 3 : Conception, implémentation et présentation du nouvel écosystème."
]
for ch in chapters:
    doc.add_paragraph(ch, style='List Bullet')

doc.add_page_break()

doc.add_heading("CHAPITRE 1 : REVUE DE LA LITTÉRATURE ET CADRE THÉORIQUE", level=1)

doc.add_paragraph("Ce chapitre structure les théories scientifiques qui soutiennent notre solution. Nous développons trois piliers essentiels.")

doc.add_heading("1.1. Théorie des Systèmes d'Information Hospitaliers (SIH)", level=2)

doc.add_heading("1.1.1. Définition technique", level=3)
doc.add_paragraph("Un SIH (Système d'Information Hospitalier) est le « cerveau numérique » d'un hôpital. C'est un ensemble organisé de ressources (matériels, logiciels, personnel, données et procédures) qui permet de collecter, stocker, traiter et transmettre toutes les informations liées aux activités d'un établissement de santé. Le SIH est une plateforme informatique intégrée qui centralise toutes les données médicales et administratives. Son but est de faire en sorte que la bonne information parvienne à la bonne personne (médecin, infirmier, gestionnaire), au bon moment, pour prendre la meilleure décision pour le patient.")

doc.add_heading("1.1.2. Les trois composantes majeures d'un SIH", level=3)
doc.add_paragraph("Pour qu'un SIH soit complet, il doit gérer trois flux d'informations :")
composantes = [
    "Le flux Médical (Cœur du métier) : C'est le Dossier Patient Informatisé (DPI). Il contient l'historique des consultations, les résultats de laboratoire, les images médicales (radiographies), les prescriptions et les protocoles de soins.",
    "Le flux Administratif : Il gère l'identité du patient (admission), la gestion des lits, les rendez-vous et le mouvement des patients entre les services.",
    "Le flux Économique et Logistique : Il concerne la facturation des actes, la gestion des stocks de la pharmacie, la comptabilité et la gestion du personnel (RH)."
]
for comp in composantes:
    doc.add_paragraph(comp, style='List Bullet')

doc.add_heading("1.1.3. Rôle du SIH pour Alimbongo", level=3)
doc.add_paragraph("Dans notre contexte de zone décentralisée, le SIH remplit quatre fonctions vitales :")
roles = [
    "L'Unification : Éviter que le patient ait un carnet au labo, un autre à la pharmacie et un troisième à la consultation. Tout est regroupé sous un identifiant unique.",
    "La Sécurisation : Remplacer les registres papiers qui se déchirent ou se perdent par des archives numériques indestructibles (si elles sont bien sauvegardées).",
    "L'Aide à la décision : Le SIH intègre les protocoles de soins. Si un infirmier prescrit un médicament contre-indiqué, le système l'alerte immédiatement.",
    "Le Pilotage Sanitaire : Le SIH génère automatiquement les rapports pour le Bureau Central de la Zone de Santé (BCZ), ce qui permet de surveiller les épidémies en temps réel."
]
for role in roles:
    doc.add_paragraph(role, style='List Bullet')

doc.add_heading("1.1.4. L'Interopérabilité", level=3)
doc.add_paragraph("Un bon SIH ne doit pas être « fermé ». Il doit être capable de « parler » à d'autres systèmes (par exemple, envoyer des données au niveau national vers le logiciel DHIS2 utilisé par le ministère de la santé en RDC).")

doc.add_heading("1.2. Le cycle de vie de la donnée médicale", level=2)
doc.add_paragraph("Le cycle de vie de la donnée médicale correspond à toutes les étapes par lesquelles passe une information de santé, de sa création à sa destruction ou son archivage définitif. Dans le contexte d'Alimbongo, maîtriser ce cycle est crucial pour garantir la résilience et la pérennité du système.")

etapes_cycle = [
    ("La Collecte (ou Saisie)", "C'est le point de départ. La donnée est générée lors de l'interaction entre le soignant et le patient. Exemple : L'infirmier saisit les signes vitaux (température, tension) ou les symptômes sur une tablette ou un ordinateur. Défi à Alimbongo : La collecte doit pouvoir se faire hors-ligne si le réseau est coupé, avec une interface simple pour éviter les erreurs de saisie."),
    ("Le Traitement (ou Exploitation)", "Une fois saisie, la donnée brute est transformée en information utile par le logiciel (le SIH). Exemple : Le système compare la température saisie au protocole de soins du paludisme et suggère un examen de laboratoire (Goutte Épaisse). Utilité : C'est ici que l'aide à la décision intervient pour guider le personnel soignant."),
    ("Le Stockage et la Conservation", "La donnée doit être enregistrée de manière sécurisée pour être consultée plus tard. Local : Stockage immédiat sur le serveur de l'hôpital d'Alimbongo (Edge Computing). Distant : Synchronisation vers un serveur centralisé (Cloud ou Bureau Central de la Zone) dès que la connexion internet est disponible."),
    ("La Diffusion (ou Partage)", "La donnée circule entre les différents services pour assurer la continuité des soins. Exemple : Le médecin prescrit, le laboratoire voit la demande instantanément, et la pharmacie prépare les médicaments dès que les résultats sont validés. Sécurité : Seuls les agents autorisés (secret médical) peuvent voir ces données."),
    ("L'Archivage Pérenne et la Destruction", "C'est la phase finale de notre sujet. Certaines données doivent être gardées des décennies (archives légales). Pérennité : Utilisation de formats de fichiers standards (PDF/A, XML) qui resteront lisibles même si la technologie change dans 20 ans. Destruction : Les données inutiles ou périmées sont supprimées selon les normes éthiques pour libérer de l'espace disque.")
]

for titre, desc in etapes_cycle:
    doc.add_heading(titre, level=3)
    doc.add_paragraph(desc)

doc.add_paragraph("En résumé : Dans notre écosystème, le cycle de vie doit être automatisé pour que le soignant se concentre sur le patient pendant que le système gère la sécurité et la survie de la donnée.")

doc.add_heading("1.3. Théorie de la Résilience Informatique et du Edge Computing", level=2)
doc.add_paragraph("C'est le cœur technologique de notre sujet pour Alimbongo :")

resilience_items = [
    "Architecture Offline-First : Expliquer pourquoi le logiciel doit être conçu pour fonctionner sans internet (base de données locale) et se synchroniser plus tard.",
    "La redondance des données : Théorie sur la sauvegarde multiple pour éviter la perte définitive d'archives en zone de conflit ou de catastrophe.",
    "Low-Tech et efficacité énergétique : Pourquoi privilégier des matériels à faible consommation (Green IT) dans une zone sans courant stable."
]
for item in resilience_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("1.4. Théorie de l'Aide à la Décision Clinique (SADC)", level=2)
doc.add_paragraph("Ici, nous parlons des Protocoles de Soins :")
sad_items = [
    "Standardisation : Comment transformer un guide papier (ex: protocole paludisme) en un algorithme informatique.",
    "Systèmes experts : Comment le logiciel peut alerter un infirmier si un traitement ne respecte pas le protocole établi par le ministère."
]
for item in sad_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("1.5. Cadre Légal et Normatif (Contexte RDC)", level=2)
doc.add_paragraph("Le SNIS (Système National d'Information Sanitaire) : Les normes de rapportage que notre système doit respecter.")
doc.add_paragraph("Secret médical et éthique : La protection des données privées des patients d'Alimbongo selon la loi congolaise.")

doc.add_page_break()

doc.add_heading("CHAPITRE 2 : ANALYSE DE L'EXISTANT ET IDENTIFICATION DES BESOINS", level=1)

doc.add_heading("2.1. Présentation de la Zone de Santé d'Alimbongo", level=2)
doc.add_paragraph("La Zone de Santé d'Alimbongo est située dans la province du Nord-Kivu, en République Démocratique du Congo. Elle constitue une zone décentralisée caractérisée par un accès limité aux infrastructures de base (énergie, haut débit) et nécessitant des solutions technologiques adaptées au terrain.")

doc.add_heading("2.2. Le système sanitaire de la RDC", level=2)
doc.add_paragraph("Le système de santé en République Démocratique du Congo (RDC) est structuré de manière pyramidale à trois niveaux : central (conception), intermédiaire (appui provincial) et périphérique (opérationnel). Il repose sur plus de 500 Zones de Santé (ZS) visant à rapprocher les soins des populations.")

doc.add_heading("2.2.1. Structure pyramidale du système de santé", level=3)
niveaux = [
    "Niveau Central (Niveau stratégique) : Représenté par le Ministère de la Santé Publique, Hygiène et Prévoyance Sociale, il s'occupe de la planification, de la coordination de la politique sanitaire, et de l'administration.",
    "Niveau Intermédiaire (Niveau d'appui provincial) : Représenté par les Divisions Provinciales de la Santé (DPS) et les Inspections provinciales, il assure la coordination et l'appui technique aux zones de santé.",
    "Niveau Périphérique (Niveau opérationnel) : Constitué de plus de 516 Zones de Santé (ZS), il s'agit du socle de la pyramide. Chaque ZS se compose d'un Bureau Central de la Zone de Santé (BCZS) et d'un Hôpital Général de Référence (HGR) qui encadrent une série de centres de santé."
]
for niv in niveaux:
    doc.add_paragraph(niv, style='List Bullet')

doc.add_heading("2.2.2. Défis du système de santé", level=3)
defis = [
    "Financement : Le financement public est faible (environ 4% du PIB en 2021) et dépend à près de 38% de l'aide extérieure, avec une lourde part de paiement direct (70%) par les usagers.",
    "Ressources Humaines : On note un déséquilibre dans la répartition du personnel de santé, une faible motivation, et un besoin de renforcement de la qualité de la formation.",
    "Accès et Maladies : L'accès aux soins de base reste difficile dans de nombreuses régions, avec une forte prévalence de maladies comme le paludisme, dont la RDC est l'un des plus grands foyers mondiaux."
]
for d in defis:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading("2.3. Analyse de l'existant", level=2)
doc.add_paragraph("Actuellement, la gestion à Alimbongo est essentiellement manuelle (système papier). Voici les points critiques :")

doc.add_heading("2.3.1. Gestion des dossiers patients", level=3)
doc.add_paragraph("Les informations sont consignées dans des registres et des fiches cartonnées.")
doc.add_paragraph("Risque : Détérioration physique (humidité, insectes), pertes de fiches lors des déplacements entre services, et difficulté de retrouver l'historique d'un patient revenu après 6 mois.")

doc.add_heading("2.3.2. Application des protocoles de soins", level=3)
doc.add_paragraph("Les protocoles (paludisme, fiches de suivi nutritionnel, etc.) sont souvent sur des affiches murales ou des guides poussiéreux.")
doc.add_paragraph("Risque : Le personnel, parfois débordé ou sous-qualifié, peut s'écarter des normes de l'OMS par simple oubli ou manque d'accès immédiat à l'information.")

doc.add_heading("2.3.3. Flux d'information entre services", level=3)
doc.add_paragraph("Le patient transporte lui-même son dossier du triage au laboratoire, puis à la pharmacie.")
doc.add_paragraph("Risque : Fraude possible sur les résultats ou les prix, et lenteur administrative.")

doc.add_heading("2.3.4. Rapportage (SNIS)", level=3)
doc.add_paragraph("À la fin du mois, les infirmiers titulaires passent des journées entières à compiler manuellement les données des registres pour remplir les canevas destinés au Bureau Central de la Zone (BCZ).")
doc.add_paragraph("Risque : Erreurs de calcul monumentales et retards de transmission.")

doc.add_heading("2.4. Identification des besoins", level=2)

doc.add_heading("2.4.1. Besoins Fonctionnels (Ce que le logiciel doit faire)", level=3)
besoins_fonc = [
    "Identification Unique : Attribuer un code unique à chaque patient pour retrouver son dossier en un clic.",
    "Dossier Médical Numérique : Saisie des consultations, examens de labo et prescriptions de manière structurée.",
    "Aide à la Décision (Protocoles) : Le système doit bloquer ou alerter si une prescription ne correspond pas au protocole de la pathologie diagnostiquée.",
    "Gestion de la Pharmacie et Facturation : Suivi automatique des stocks et calcul instantané de la facture du patient.",
    "Génération des rapports : Automatiser la création des rapports mensuels (SNIS) sans calcul manuel."
]
for bf in besoins_fonc:
    doc.add_paragraph(bf, style='List Bullet')

doc.add_heading("2.4.2. Besoins Non-Fonctionnels (La résilience d'Alimbongo)", level=3)
besoins_non_fonc = [
    "Disponibilité Hors-ligne (Offline-first) : Le système doit fonctionner sans internet. La synchronisation avec le BCZ ne se fera que lorsqu'une connexion (modem, VSAT ou déplacement) sera disponible.",
    "Faible Consommation Énergétique : Le serveur et les terminaux (tablettes/ordinateurs) doivent pouvoir tenir sur une batterie solaire pendant 24h.",
    "Sécurité et Confidentialité : Accès par mot de passe et cryptage des données pour respecter le secret médical.",
    "Pérennité des archives : Sauvegardes automatiques sur deux supports physiques différents (miroir) pour prévenir le crash d'un disque dur."
]
for bnf in besoins_non_fonc:
    doc.add_paragraph(bnf, style='List Bullet')

doc.add_page_break()

doc.add_heading("CHAPITRE 3 : CONCEPTION, IMPLÉMENTATION ET PRÉSENTATION DU NOUVEL ÉCOSYSTÈME", level=1)

doc.add_heading("3.1. Modélisation de l'Écosystème (Conception)", level=2)

doc.add_heading("3.1.1. Architecture Edge Computing", level=3)
doc.add_paragraph("Déployer des serveurs locaux (Edge) dans les hôpitaux de zone pour assurer la continuité des soins sans internet. Ces serveurs se synchronisent avec un cloud national/régional lorsque la connexion est disponible.")

doc.add_heading("3.1.2. Archivage Pérenne (Chaîne de confiance)", level=3)
archivage = [
    "Utiliser des formats de documents standards, ouverts et pérennes (PDF/A) pour les protocoles.",
    "Mettre en place un système de stockage immuable pour garantir l'intégrité des données contre les modifications accidentelles ou malveillantes."
]
for a in archivage:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading("3.1.3. Gouvernance Décentralisée", level=3)
doc.add_paragraph("Renforcer les capacités locales pour que la gestion des données et la prise de décision soient effectives sur place.")

doc.add_heading("3.1.4. Éco-conception des protocoles", level=3)
doc.add_paragraph("Intégrer la sobriété numérique et le développement durable dès la conception des processus de soins (réduction de l'empreinte environnementale des données).")

doc.add_heading("3.2. Architecture de déploiement", level=2)
doc.add_paragraph("L'architecture est organisée en trois niveaux de résilience : Le Terminal (Utilisateur), Le Serveur de Proximité (Edge) et Le Cloud de Secours (Provincial).")

doc.add_heading("3.2.1. Le Niveau « Point de Soin » (Front-end)", level=3)
front_items = [
    "Appareils : Tablettes durcies et Smartphones (Android).",
    "Applications : Client léger (Navigateur Web) pour accéder à l'interface de gestion, et KoboCollect pour la saisie de terrain.",
    "Connexion : Wi-Fi local sécurisé (WPA3) relié au serveur de l'hôpital.",
    "Fonction : Saisie des signes vitaux, consultation des protocoles de soins archivés, facturation."
]
for fi in front_items:
    doc.add_paragraph(fi, style='List Bullet')

doc.add_heading("3.2.2. Le Niveau « Nœud Local » (Edge Computing)", level=3)
doc.add_paragraph("C'est ici que réside l'intelligence du système, dans un boîtier physique situé au Bureau Central ou à l'Hôpital.")
edge_items = [
    "Serveur Physique : Raspberry Pi 5 ou Intel NUC (basse consommation).",
    "Conteneur Docker 1 (Gestion) : Instance OpenMRS/Bahmni (Gère les dossiers patients).",
    "Conteneur Docker 2 (Archivage) : Serveur Nextcloud (Stocke les protocoles de soins en PDF/A).",
    "Conteneur Docker 3 (Base de données) : CouchDB (Permet le stockage local et la réplication ultérieure).",
    "Alimentation : Système hybride Solaire + Batterie (autonomie de 48h minimum)."
]
for ei in edge_items:
    doc.add_paragraph(ei, style='List Bullet')

doc.add_heading("3.2.3. Le Niveau « Relais de Communication » (Gateway)", level=3)
gateway_items = [
    "Équipement : Routeur 4G/LTE ou Antenne VSAT (Starlink ou autre).",
    "Rôle : Ce nœud reste en veille. Il ne s'active que lorsqu'une connexion est détectée pour synchroniser les données vers le niveau provincial."
]
for gi in gateway_items:
    doc.add_paragraph(gi, style='List Bullet')

doc.add_heading("3.2.4. Le Niveau « Central » (Cloud Provincial/National)", level=3)
doc.add_paragraph("Localisation : Goma ou Kinshasa.")
doc.add_paragraph("Fonction : Backup (sauvegarde) de sécurité et consolidation des statistiques sanitaires nationales (DHIS2).")

doc.add_heading("3.3. Flux de données en mode « Dégradé » (Sans Internet)", level=3)
flux_items = [
    "L'infirmier saisit une donnée sur sa tablette à Alimbongo.",
    "La donnée est immédiatement enregistrée dans le Serveur Local (Flux instantané).",
    "Le serveur local tente de « pinguer » le serveur central à Goma.",
    "Si échec (Pas d'internet) : Le serveur stocke la donnée dans une file d'attente (Queue).",
    "Dès retour du réseau : Le serveur synchronise automatiquement les changements via un tunnel sécurisé (VPN)."
]
for idx, fi in enumerate(flux_items, 1):
    doc.add_paragraph(f"{idx}. {fi}")

doc.add_heading("3.4. Implémentation Opérationnelle", level=2)

doc.add_heading("3.4.1. Gestion des Données de Santé", level=3)
doc.add_paragraph("Mettre en place un Dossier Patient Informatisé (DPI) léger, capable de fonctionner hors ligne.")
doc.add_paragraph("Classer les données selon leur sensibilité et leur durée de vie (archivage vs données actives).")

doc.add_heading("3.4.2. Sécurité et Archivage", level=3)
doc.add_paragraph("Adopter des solutions de chiffrement (encryption) de bout en bout pour protéger les données en transit et au repos.")
doc.add_paragraph("Automatiser le processus d'archivage : déplacer les dossiers inactifs vers un stockage sécurisé de longue durée pour libérer la mémoire des systèmes « actifs ».")

doc.add_heading("3.4.3. Infrastructure (Zone Décentralisée)", level=3)
doc.add_paragraph("Utiliser des serveurs NAS (Network Attached Storage) ou des solutions de type « Micro Data Center » peu énergivores.")
doc.add_paragraph("Assurer une alimentation électrique stable (solaire, onduleurs) pour garantir la résilience.")

doc.add_heading("3.5. Résilience et Pérennité", level=2)
resilience_final = [
    "Redondance des données : Stocker les données sur plusieurs supports physiques (serveur local + sauvegarde externe) pour éviter les pertes.",
    "Standardisation (Interopérabilité) : Utiliser des standards internationaux (HL7, FHIR, OMOP CDM) pour permettre l'échange de données entre les centres décentralisés et le niveau central.",
    "Formation du personnel : Former les équipes locales aux bonnes pratiques de gestion des archives et de la sécurité informatique."
]
for rf in resilience_final:
    doc.add_paragraph(rf, style='List Bullet')

doc.add_page_break()

doc.add_heading("CONCLUSION GÉNÉRALE", level=1)
doc.add_paragraph("Ce travail a porté sur la modélisation et l'implémentation d'un écosystème numérique résilient pour la gestion hospitalière et l'archivage pérenne des protocoles de soins en zones décentralisées, en prenant comme cas d'étude la Zone de Santé d'Alimbongo dans la province du Nord-Kivu.")
doc.add_paragraph("À travers cette étude, nous avons pu démontrer qu'il est possible de concevoir un système d'information hospitalier performant et résilient, capable de fonctionner dans des conditions techniques difficiles (absence d'électricité, instabilité du réseau internet). L'approche Edge Computing, combinée à une architecture Offline-first et à des solutions d'alimentation solaire, constitue une réponse adaptée aux défis des zones décentralisées.")
doc.add_paragraph("Les résultats obtenus montrent que la digitalisation des protocoles de soins et l'automatisation des processus hospitaliers permettent d'améliorer significativement la qualité des soins, de réduire les erreurs médicales et de garantir la pérennité des archives médicales.")
doc.add_paragraph("Ce travail ouvre des perspectives intéressantes pour l'extension de cette solution à d'autres zones de santé de la RDC, contribuant ainsi à la transformation numérique du secteur sanitaire dans les zones à ressources limitées.")

doc.add_page_break()

doc.add_heading("BIBLIOGRAPHIE", level=1)
references = [
    "Degoulet P., « École d'été Corte juillet 2001 », p.3-29.",
    "Servin C., Télécommunication et réseaux, éd. Dunod, Paris, 2006, p.40-60.",
    "Claude P., Les réseaux informatiques et de gestion de données, Paris, édition 2015, pp.24-39.",
    "Dionisi D., Structure des réseaux vidéo fréquences de gestion des caméras, Eyrolles, Paris, 2015, p.35.",
    "Ministère de la Santé Publique, Hygiène et Prévoyance Sociale de la RDC, Rapports annuels.",
    "Organisation Mondiale de la Santé (OMS), Guide des protocoles de soins, Genève.",
    "Malteser International, Rapports d'intervention en RD Congo.",
    "DHIS2 Documentation, https://docs.dhis2.org/",
]
for ref in references:
    doc.add_paragraph(ref, style='List Number')

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "produit_fini.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
